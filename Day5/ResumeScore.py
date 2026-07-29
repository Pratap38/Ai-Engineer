from pydantic import BaseModel,Field
import os
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
import time
load_dotenv()

# Retrieve API key
my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key missing")

# Initialize the client
client = Groq(api_key=my_api_key)

# Define model and message
model_name = "llama-3.3-70b-versatile"
jobDescription="""
Job Description
Wissen Technology is hiring for Python Developer

About Wissen Technology: 

At Wissen Technology, we deliver niche, custom-built products that solve complex business challenges across industries worldwide. Founded in 2015, our core philosophy is built around a strong product engineering mindset—ensuring every solution is architected and delivered right the first time. Today, Wissen Technology has a global footprint with 2000+ employees across offices in the US, UK, UAE, India, and Australia. Our commitment to excellence translates into delivering 2X impact compared to traditional service providers. How do we achieve this? Through a combination of deep domain knowledge, cutting-edge technology expertise, and a relentless focus on quality. We don’t just meet expectations—we exceed them by ensuring faster time-to-market, reduced rework, and greater alignment with client objectives. We have a proven track record of building mission-critical systems across industries, including financial services, healthcare, retail, manufacturing, and more. Wissen stands apart through its unique delivery models. Our outcome-based projects ensure predictable costs and timelines, while our agile pods provide clients with the flexibility to adapt to their evolving business needs. Wissen leverages its thought leadership and technology prowess to drive superior business outcomes. Our success is powered by top-tier talent.  Our mission is clear: to be the partner of choice for building world-class custom products that deliver exceptional impact—the first time, every time.

Job Summary: 

We are seeking a skilled Python Developer to design, develop, and maintain scalable applications and backend services. The ideal candidate will have strong expertise in Python, problem-solving, API development, and database technologies, with the ability to deliver high-quality, efficient, and reliable software solutions.

Experience: 2-6 Years 
Location: Mumbai
Mode of Work: Full Time
Education: B.Tech/M.Tech in Computer Science or related field
Key Responsibilities:

Experience in Python (Only Backend), Data structures, Oops, Algorithms, Django, NumPy etc. 
Notice/Joining of not more than 30 days. 
Hybrid Mode of working. 
Good understanding of writing Unit Tests using PYTest. 
Good understanding of parsing XML’s and handling files using Python. 
Good understanding with Databases/SQL, procedures and query tuning. 
Service Design Concepts, OO and Functional Development concepts. 
Agile Development Methodologies. 
Strong oral and written communication skills. 
Excellent interpersonal skills and professional approach Skills desired
Wissen Sites: 

Website: www.wissen.com
LinkedIn: https://www.linkedin.com/company/wissen-technology
Wissen Leadership: Leadership Team | Wissen
Wissen Live: Log Masuk LinkedIn, Daftar Masuk | LinkedIn
Wissen Thought Leadership: https://www.wissen.com/articles/
"""

class JobDesc(BaseModel):
    role:str
    requiredSkill:list[str]
    preferSkill:list[str]
    minExperience:float|None
    educationRequirement:list[str]
    responsibility:list[str]

jobSchema=JobDesc.model_json_schema()

system_prompt=f"""
You are an expert HR assistant 

your main goal is to analyse this resume and extract the structure information from it

Return only the valid json Schema that matches this {jobSchema}
Important:
Donot return the schema itself
Donot return any useless info like the title summary etc.

if any experiecne is not then return null
if any requred info is missing then just simply return an empty list for that field
donot add anything else by own just need as per the resume and schema.
"""

userPrompt=f"""
Analyse the following  job description
{jobDescription}
"""
messageSystem={
    "role":"system",
    "content":system_prompt
}
messageUser={
    "role":"user",
    "content":userPrompt
}
responseformat={
    "type":"json_object"
}
messages=[messageSystem,messageUser]
response = client.chat.completions.create(
    model=model_name,
    messages=messages,
    response_format=responseformat,

)
answer=response.choices[0].message.content

rawJson=answer

import json

jsonData=json.loads(rawJson)

job=JobDesc(**jsonData)


class MatchResult(BaseModel):
    score:float
    details:dict

class ExperienceResume(BaseModel):
    company:str|None=None
    role:str|None=None
    duration:str|None=None
    skillUsed:list[str]=[]

class Resume(BaseModel):
    name:str|None=None
    email:str|None=None
    phone:str|None=None
    totalExperienceinYears:float|None=None
    skills:list[str]=[]
    experience:list[ExperienceResume]=[]
    project:list[str]=[]
    certification:list[str]=[]

resumeSchema=Resume.model_json_schema()

def finalScore(job,resume):
    matchschema=MatchResult.model_json_schema()
    prompt=f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {matchschema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model_name, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def paresedresume(resumetext):
    system_prompt=f"""
You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resumeSchema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resumetext}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model_name, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume




##resume read method
from PyPDF2 import PdfReader
from docx import Document

def Readpdf(filePath):
    reader=PdfReader(filePath)
    text=""
    for page in reader.pages:
        pagetext=page.extract_text()
        if pagetext:
            text+=pagetext+"\n"
    return text

def Readdoc(filePath):
    reader=Document(filePath)
    text=""
    for para in reader.paragraphs:
        if para.text.strip():
            text+=para.text+"\n"
    for table in reader.tables:
        for row in table.rows:
            for cell in row.cells:        ##for coloum jo ki resume me present and upar wala row mean kch resume me table bane rehte
                if cell.text.strip():
                    text+=cell.text+"\n"
    return text

def readResume(filePath):
    if filePath.suffix.lower()==".pdf":
        return Readpdf(filePath)
    elif filePath.suffix.lower()==".docx":
        return Readdoc(filePath)
    else :
        return None


resumeFolder=Path("resume")
allResult=[]
for filepath in  resumeFolder.iterdir():
    if (filepath.suffix.lower() not in [".pdf",".docx"]):
        continue
    print("processing\n",filepath.name)
    resumeText=readResume(filepath)
    parsedresume=paresedresume(resumeText)
    time.sleep(5)
    result=finalScore(job,parsedresume)
    time.sleep(5)
    print("score",result.score)
    allResult.append({
        "name":parsedresume.name,
        "score":result.score,
        "details":result.details

    })
allResult.sort(
    key=lambda candidate:candidate["score"],
    reverse=True

)
top_2 = allResult[:2]
worst_2 = allResult[-2:]


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])


